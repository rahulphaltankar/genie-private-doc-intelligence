import os
import uuid
from typing import List, Dict, Any
from pypdf import PdfReader
from docx import Document
from genie.pipeline.chunker import chunk_by_page_text
from genie.schemas.metadata_schema_v3 import ChunkMeta

def ingest_document(file_path: str) -> List[ChunkMeta]:
    """
    Ingests a document, extracts text per page/paragraph, and returns a list of ChunkMeta objects.
    """
    filename = os.path.basename(file_path)
    doc_id = str(uuid.uuid4())
    
    file_ext = os.path.splitext(filename)[1].lower()
    page_texts = []
    
    if file_ext == '.pdf':
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_texts.append(page.extract_text() or "")
            
    elif file_ext == '.docx':
        doc = Document(file_path)
        # For Word, we treat the whole doc as one "page" for now, or could split by section
        # Improved: Split by paragraphs and group into pseudo-pages or just one stream
        full_text = "\n".join([para.text for para in doc.paragraphs])
        page_texts = [full_text]
        
    elif file_ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            page_texts = [f.read()]
            
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")
        
    return chunk_by_page_text(doc_id, filename, page_texts)

def process_uploaded_files(uploaded_files) -> List[ChunkMeta]:
    """
    Processes a list of Streamlit uploaded files.
    """
    all_chunks = []
    for uploaded_file in uploaded_files:
        # Streamlit files are file-like objects
        filename = uploaded_file.name
        doc_id = str(uuid.uuid4())
        page_texts = []
        
        file_ext = os.path.splitext(filename)[1].lower()
        
        if file_ext == '.pdf':
            reader = PdfReader(uploaded_file)
            for page in reader.pages:
                page_texts.append(page.extract_text() or "")
        elif file_ext == '.docx':
            doc = Document(uploaded_file)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            page_texts = [full_text]
        elif file_ext == '.txt':
            page_texts = [str(uploaded_file.read(), "utf-8")]
            
        all_chunks.extend(chunk_by_page_text(doc_id, filename, page_texts))
        
    return all_chunks
